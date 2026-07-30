#!/usr/bin/env python3
"""
actualizar_dashboard.py
Secretaría de Ambiente · San Miguel de Tucumán
===============================================
Procesa archivos nuevos y sube todo el dashboard a GitHub automáticamente.

USO:
  python actualizar_dashboard.py                    # procesa carpeta 'nuevas_ordenes'
  python actualizar_dashboard.py ruta/archivo.docx  # procesa un archivo específico

REQUISITOS:
  pip install python-docx requests python-docx
"""

import os, sys, json, re, glob, base64, hashlib
from pathlib import Path
from datetime import datetime

# ══════════════════════════════════════════════════════════
# CONFIGURACIÓN — editá estos valores una sola vez
# ══════════════════════════════════════════════════════════
# Token — se lee desde el archivo "token.txt" en la misma carpeta
# Así no se pierde al actualizar el script
def _leer_token():
    ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)), "token.txt")
    if os.path.exists(ruta):
        with open(ruta) as f:
            t = f.read().strip()
            if t: return t
    return "ghp_XXXXXXXXXXXXXXXXXXXXXX"

GITHUB_OWNER  = "danielgucorrea-dot"
GITHUB_REPO   = "BAM-CMA"
GITHUB_BRANCH = "main"
GITHUB_TOKEN  = _leer_token()

# Carpeta donde vas a poner los .docx nuevos cada semana
CARPETA_NUEVAS = "nuevas_ordenes"

# Archivos del dashboard a subir (relativos a este script)
ARCHIVOS_DASHBOARD = {
    "pv_data.json":         "pv_data.json",
    "actas.json":           "actas.json",
    "incumplimientos.json": "incumplimientos.json",
}
# index.html y ordenes_backup.json se actualizan SOLO via API o Sync del dashboard
# ACTUALIZAR.bat NO los toca para evitar sobreescribir cambios
# ══════════════════════════════════════════════════════════

MESES = {
    'enero':1,'febrero':2,'marzo':3,'abril':4,'mayo':5,'junio':6,
    'julio':7,'agosto':8,'septiembre':9,'octubre':10,'noviembre':11,'diciembre':12
}

TIPOS_MAP = {
    'rnh':'RNH', 'residuo no habitual':'RNH',
    'basural':'Basural', 'volcadero':'Volcadero',
    'cesto':'Limpieza de cesto', 'limpieza de cesto':'Limpieza de cesto',
    'platabanda':'Platabanda', 'barrido':'Barrido',
    'poda':'Restos de poda', 'falta de recolec':'Falta de recolección',
    'falta recolec':'Falta de recolección', 'recolec':'Falta de recolección',
    'levante':'Levante de bolsas', 'escombro':'RNH', 'limpieza':'Limpieza de cesto',
}

def normalizar_tipo(raw):
    t = (raw or '').lower().strip()
    if not t or 'deficiencia detectada' in t:
        return None
    for k, v in TIPOS_MAP.items():
        if k in t:
            return v
    return 'Otro'

def cap(s):
    return s[0].upper() + s[1:] if s else s

def parsear_docx(ruta):
    """Parsea un .docx de orden de servicio → dict con key, numero, fecha, empresa, items"""
    from docx import Document

    doc   = Document(ruta)
    texto = '\n'.join(p.text for p in doc.paragraphs)
    fname = Path(ruta).stem

    # Número de orden
    m_num = re.search(r'N[ºo°][\s:]*0*(\d+)', texto, re.I)
    numero = m_num.group(1).zfill(6) if m_num else '000000'

    # Fecha desde texto "5 de mayo de 2026"
    fecha = None
    m_fecha = re.search(r'(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})', texto, re.I)
    if m_fecha:
        d, mes_str, y = m_fecha.groups()
        mes = MESES.get(mes_str.lower())
        if mes:
            fecha = f"{y}-{mes:02d}-{int(d):02d}"

    # Fallback: nombre del archivo "OrdenServ 06-05-26.docx"
    if not fecha:
        m_fn = re.search(r'(\d{2})[-_\s](\d{2})[-_\s](\d{2,4})$', fname)
        if m_fn:
            d, m, y = m_fn.groups()
            if len(y) == 2:
                y = '20' + y
            fecha = f"{y}-{m}-{d}"

    if not fecha:
        fecha = datetime.today().strftime('%Y-%m-%d')

    # Empresa
    m_emp = re.search(r'EMPRESA[\s:]+([^\n]{4,})', texto, re.I)
    empresa = m_emp.group(1).strip().split('.')[0][:60] if m_emp else 'Sin datos'

    # Items desde tablas
    items = []
    seen  = set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if len(celdas) < 3:
                continue
            n_str, tipo_raw, ubic = celdas[0], celdas[1], celdas[2]
            if not re.match(r'^\d+$', n_str):
                continue
            n    = int(n_str)
            tipo = normalizar_tipo(tipo_raw)
            if not tipo or not ubic or len(ubic) < 3:
                continue
            uid = f"{n}:{tipo}:{ubic[:20]}"
            if uid in seen:
                continue
            seen.add(uid)
            items.append({'n': n, 'tipo': tipo, 'ubic': cap(ubic)})

    items.sort(key=lambda x: x['n'])

    # Key único basado en nombre de archivo
    key = re.sub(r'[^a-zA-Z0-9_-]', '_', fname)

    print(f"  ✓ Parseado: {fname} → {fecha} | N°{numero} | {len(items)} ítems")
    return {'key': key, 'numero': numero, 'fecha': fecha, 'empresa': empresa, 'items': items}


def cargar_backup(ruta='ordenes_backup.json'):
    """Carga el backup existente o retorna estructura vacía"""
    if os.path.exists(ruta):
        with open(ruta, encoding='utf-8') as f:
            data = json.load(f)
        print(f"  Backup cargado: {len(data.get('orders',[]))} órdenes existentes")
        return data
    return {'version': 3, 'orders': [], 'mapData': [], 'mapInactivos': [], 'incumplimientos': []}


def guardar_backup(data, ruta='ordenes_backup.json'):
    with open(ruta, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, separators=(',', ':'))
    size = os.path.getsize(ruta)
    print(f"  Backup guardado: {len(data['orders'])} órdenes | {size//1024} KB")


def subir_a_github(ruta_local, ruta_repo):
    """Sube un archivo a GitHub via API"""
    try:
        import urllib.request, urllib.error
    except ImportError:
        print("  ⚠ urllib no disponible")
        return False

    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{ruta_repo}"
    headers = {
        'Authorization': f'token {GITHUB_TOKEN}',
        'Accept': 'application/vnd.github.v3+json',
        'Content-Type': 'application/json',
    }

    # Leer archivo
    with open(ruta_local, 'rb') as f:
        contenido = f.read()
    contenido_b64 = base64.b64encode(contenido).decode()

    # Obtener SHA actual (necesario para actualizar)
    sha = None
    try:
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req) as resp:
            data = json.loads(resp.read())
            sha = data.get('sha')
    except urllib.error.HTTPError as e:
        if e.code != 404:
            print(f"  ⚠ Error obteniendo SHA de {ruta_repo}: {e}")

    # Preparar payload
    fecha_str = datetime.now().strftime('%Y-%m-%d %H:%M')
    payload = {
        'message': f'Actualización automática {fecha_str} — {ruta_repo}',
        'content': contenido_b64,
        'branch':  GITHUB_BRANCH,
    }
    if sha:
        payload['sha'] = sha

    # Subir
    try:
        data_bytes = json.dumps(payload).encode()
        req = urllib.request.Request(url, data=data_bytes, headers=headers, method='PUT')
        with urllib.request.urlopen(req) as resp:
            result = json.loads(resp.read())
            accion = 'Actualizado' if sha else 'Creado'
            size_kb = len(contenido) // 1024
            print(f"  ✓ {accion}: {ruta_repo} ({size_kb} KB)")
            return True
    except urllib.error.HTTPError as e:
        error_body = e.read().decode()
        print(f"  ✗ Error subiendo {ruta_repo}: HTTP {e.code} — {error_body[:200]}")
        return False
    except Exception as e:
        print(f"  ✗ Error subiendo {ruta_repo}: {e}")
        return False


NUMERALES_INC = ['3.4.2-01','3.4.2-03','3.4.2-08','3.4.3-03','3.4.3-05']
MESES_INC = {'enero':1,'febrero':2,'marzo':3,'abril':4,'mayo':5,'junio':6,
             'julio':7,'agosto':8,'septiembre':9,'octubre':10,'noviembre':11,'diciembre':12}

def parse_val_inc(v):
    v = str(v).strip().replace('-.-','').replace('S/Nov.','').strip()
    try: return int(v)
    except: return 0

def parse_fecha_inc(raw, titulo):
    raw = raw.strip()
    mes_tit = next((v for k,v in MESES_INC.items() if k in titulo.lower()), None)
    m_full = re.match(r'(\d{1,2})/(\d{1,2})/(\d{2,4})', raw)
    if m_full:
        d, mo, y = int(m_full.group(1)), int(m_full.group(2)), int(m_full.group(3))
        if y < 100: y += 2000
        return f"{y}-{mo:02d}-{d:02d}"
    m_day = re.match(r'^(\d{1,2})$', raw)
    if m_day and mes_tit:
        return f"2026-{mes_tit:02d}-{int(m_day.group(1)):02d}"
    return None

def parsear_incumplimientos(ruta):
    """Parsea un .docx de incumplimientos → lista de registros"""
    from docx import Document
    doc = Document(ruta)
    registros = []
    for tabla in doc.tables:
        titulo = tabla.rows[0].cells[0].text.strip()
        for row in tabla.rows[2:]:
            cells = [c.text.strip() for c in row.cells]
            fecha = parse_fecha_inc(cells[0], titulo)
            if not fecha: continue
            if not fecha.startswith('202'): continue
            nums = {NUMERALES_INC[i]: parse_val_inc(cells[3+i]) for i in range(5) if 3+i < len(cells)}
            registros.append({
                'fecha':    fecha,
                'quincena': titulo.split('–')[0].strip(),
                'orden':    cells[1] if len(cells)>1 else '',
                'acta':     cells[2] if len(cells)>2 else '',
                'obs':      cells[8] if len(cells)>8 else '',
                'total':    sum(nums.values()),
                **nums
            })
    registros.sort(key=lambda r: r['fecha'])
    print(f"  ✓ Incumplimientos: {len(registros)} registros")
    return registros

def es_incumplimiento(nombre):
    """Detecta si un archivo es de incumplimientos por su nombre"""
    n = nombre.lower()
    return any(k in n for k in ['incumpl', 'numeral', 'quincenal'])

def main():
    print("=" * 55)
    print("  Dashboard Ambiental SMT — Actualizador automático")
    print("=" * 55)

    # Descargar index.html actualizado de GitHub antes de cualquier operación
    # Esto evita que el archivo local viejo sobreescriba cambios remotos
    try:
        import urllib.request as _ur
        _url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/index.html"
        _req = _ur.Request(_url, headers={'Authorization':f'token {GITHUB_TOKEN}','Accept':'application/vnd.github.v3+json'})
        with _ur.urlopen(_req) as _resp:
            import base64 as _b64, json as _js
            _data = _js.loads(_resp.read())
            _content = _b64.b64decode(_data['content'])
            with open('index.html','wb') as _f:
                _f.write(_content)
        print(f"  ✓ index.html sincronizado desde GitHub ({len(_content)//1024} KB)")
    except Exception as _e:
        print(f"  ⚠ No se pudo sincronizar index.html: {_e}")

    # Verificar token configurado
    if 'XXXX' in GITHUB_TOKEN:
        print("\n⚠ ATENCIÓN: Falta el archivo token.txt")
        print("  Creá el archivo C:\\Dashboard_SMT\\token.txt con tu token de GitHub\n")

    # Determinar archivos .docx a procesar
    docx_files = []
    if len(sys.argv) > 1:
        # Archivo(s) pasados como argumento
        for arg in sys.argv[1:]:
            if os.path.exists(arg) and arg.endswith('.docx'):
                docx_files.append(arg)
            else:
                print(f"  ⚠ No encontrado o no es .docx: {arg}")
    else:
        # Buscar en carpeta 'nuevas_ordenes'
        os.makedirs(CARPETA_NUEVAS, exist_ok=True)
        docx_files = glob.glob(os.path.join(CARPETA_NUEVAS, '*.docx'))
        if not docx_files:
            print(f"\n📁 No hay archivos .docx nuevos en '{CARPETA_NUEVAS}/'")
            print(f"   Subiendo archivos del dashboard a GitHub...")
            if 'XXXX' not in GITHUB_TOKEN:
                exitos = 0
                for local, repo in ARCHIVOS_DASHBOARD.items():
                    if os.path.exists(local):
                        ok = subir_a_github(local, repo)
                        if ok: exitos += 1
                print(f"\n{'='*55}")
                if exitos:
                    print(f"✅ {exitos} archivos subidos correctamente.")
                    print(f"   https://{GITHUB_OWNER}.github.io/{GITHUB_REPO}/")
                else:
                    print(f"❌ No se pudo subir. Revisá el token.txt")
                print(f"{'='*55}")
            else:
                print("⚠ Falta el archivo token.txt con tu token de GitHub")
            return

    # Cargar backup existente
    print(f"\n📂 Cargando backup existente...")
    backup = cargar_backup()
    keys_existentes = {o['key'] for o in backup['orders']}

    # Parsear nuevos .docx
    print(f"\n📄 Procesando {len(docx_files)} archivo(s)...")
    nuevas = 0
    duplicados = 0
    errores = 0

    inc_registros = []  # acumular incumplimientos de todos los docx

    for ruta in sorted(docx_files):
        fname = Path(ruta).name
        print(f"\n  → {fname}")
        try:
            if es_incumplimiento(fname):
                # Procesar como incumplimientos
                regs = parsear_incumplimientos(ruta)
                inc_registros.extend(regs)
                nuevas += 1
            else:
                # Procesar como orden de servicio
                orden = parsear_docx(ruta)
                if orden['key'] in keys_existentes:
                    print(f"    ⏭ Ya existe (duplicado)")
                    duplicados += 1
                else:
                    backup['orders'].append(orden)
                    keys_existentes.add(orden['key'])
                    nuevas += 1
        except Exception as e:
            print(f"    ✗ Error: {e}")
            errores += 1

    # Guardar incumplimientos si hubo nuevos
    if inc_registros:
        inc_path = 'incumplimientos.json'
        # Cargar existentes y mergear sin duplicar por fecha
        if os.path.exists(inc_path):
            with open(inc_path, encoding='utf-8') as f:
                existentes = json.load(f)
            fechas_exist = {r['fecha']+r['quincena'] for r in existentes}
            nuevos_inc = [r for r in inc_registros if r['fecha']+r['quincena'] not in fechas_exist]
            existentes.extend(nuevos_inc)
            existentes.sort(key=lambda r: r['fecha'])
            final_inc = existentes
        else:
            final_inc = inc_registros
        with open(inc_path, 'w', encoding='utf-8') as f:
            json.dump(final_inc, f, ensure_ascii=False, separators=(',',':'))
        print(f"\n✓ incumplimientos.json: {len(final_inc)} registros guardados")

    # Ordenar por fecha descendente
    backup['orders'].sort(key=lambda o: o['fecha'], reverse=True)

    print(f"\n📊 Resultado: {nuevas} nuevas | {duplicados} duplicadas | {errores} errores")

    if nuevas == 0 and not inc_registros:
        print("  No hay archivos nuevos para procesar.")
        # Igual subir los archivos del dashboard si existen localmente
        if 'XXXX' not in GITHUB_TOKEN:
            print(f"\n🚀 Subiendo archivos del dashboard a GitHub...")
            exitos = 0
            for local, repo in ARCHIVOS_DASHBOARD.items():
                if os.path.exists(local):
                    ok = subir_a_github(local, repo)
                    if ok: exitos += 1
            if exitos:
                print(f"\n✅ {exitos} archivos subidos correctamente.")
        return

    # Guardar backup local
    print(f"\n💾 Guardando backup local...")
    guardar_backup(backup)

    # Subir a GitHub
    if 'XXXX' in GITHUB_TOKEN:
        print(f"\n⚠ Token no configurado — saltando subida a GitHub.")
        print(f"  El archivo ordenes_backup.json fue actualizado localmente.")
        print(f"  Subilo manualmente o configurá el token.")
        return

    print(f"\n🚀 Subiendo archivos a GitHub ({GITHUB_OWNER}/{GITHUB_REPO})...")
    exitos = 0
    for local, repo in ARCHIVOS_DASHBOARD.items():
        if os.path.exists(local):
            ok = subir_a_github(local, repo)
            if ok:
                exitos += 1
        else:
            print(f"  ⚠ No encontrado localmente: {local} (saltando)")

    print(f"\n{'='*55}")
    if exitos > 0:
        print(f"✅ Subida completa: {exitos}/{len(ARCHIVOS_DASHBOARD)} archivos")
        print(f"   Dashboard disponible en ~2 minutos:")
        print(f"   https://{GITHUB_OWNER}.github.io/{GITHUB_REPO}/")
    else:
        print(f"❌ No se pudo subir ningún archivo. Revisá el token.")
    print(f"{'='*55}\n")

    # Mover procesados a carpeta 'procesadas'
    if docx_files and CARPETA_NUEVAS in str(docx_files[0]):
        carpeta_proc = os.path.join(CARPETA_NUEVAS, 'procesadas')
        os.makedirs(carpeta_proc, exist_ok=True)
        for ruta in docx_files:
            dest = os.path.join(carpeta_proc, Path(ruta).name)
            os.rename(ruta, dest)
        print(f"  📁 .docx movidos a '{carpeta_proc}/'")


if __name__ == '__main__':
    main()
